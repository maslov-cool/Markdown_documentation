import docx


def markdown_to_docx(text):
    lines = text.split('\n')
    doc = docx.Document()
    doc.add_heading(lines[0], level=0)
    cnt = False
    for i in lines[1:]:
        if i:
            cnt = False
            if i[:7].count('#'):
                n = i[:7].count('#')
                p = doc.add_heading(level=n)
                i = i[n + 1:]
            elif i[0].isdigit() and i[1] == '.':
                i = i[3:]
                p = doc.add_paragraph(style='List Number')
            elif i[:2] in ['* ', '- ', '+ ']:
                i = i[2:]
                p = doc.add_paragraph(style='List Bullet')
            else:
                p = doc.add_paragraph()
            c = ''
            if not i.count('*') % 2 and i.count('_') % 2 > 0:
                c = '*'
            elif i.count('*') % 2 > 0 and not i.count('_') % 2:
                c = '_'
            elif i.count('*') % 2 == i.count('_') % 2:
                c = '*_'
            if c:
                cnt = 0
                flag_italic, flag_bold = False, False
                for s in i:
                    if s in c:
                        cnt += 1
                    else:
                        if cnt == 1:
                            flag_italic = not flag_italic
                        elif cnt == 2:
                            flag_bold = not flag_bold
                        elif cnt == 3:
                            flag_bold = not flag_bold
                            flag_italic = not flag_italic
                        run = p.add_run(s)
                        run.italic = flag_italic
                        run.bold = flag_bold
                        cnt = 0
            else:
                p.add_run(i)
        else:
            if cnt:
                doc.add_paragraph()
            cnt = True
    doc.save(f'{lines[0]}.docx')

